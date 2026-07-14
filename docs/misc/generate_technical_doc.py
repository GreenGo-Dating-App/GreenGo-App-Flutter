"""
GreenGo Technical Documentation Generator - Engineering Edition
Comprehensive DOCX for the engineering team with code patterns,
architecture details, security rules, build config, and more.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row_data in enumerate(rows):
        for ci, cell_data in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(cell_data)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def add_code_block(doc, code, language=''):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    return p

def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

def numbered(doc, text):
    doc.add_paragraph(text, style='List Number')

def create_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── COVER PAGE ──
    for _ in range(3): doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('GreenGo')
    r.font.size = Pt(48); r.font.color.rgb = RGBColor(0x60,0xA9,0x17); r.bold = True

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('Technical Engineering Documentation')
    r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x33,0x33,0x33)

    doc.add_paragraph()
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = m.add_run('Version 1.1.0+26 | March 2026\nFor the Engineering Team\nConfidential')
    r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    doc.add_page_break()

    # ── TABLE OF CONTENTS ──
    doc.add_heading('Table of Contents', level=1)
    toc = [
        '1. Executive Summary & System Overview',
        '2. System Architecture & Component Diagram',
        '3. Technology Stack & Dependencies',
        '4. Frontend Architecture',
        '   4.1 Clean Architecture Pattern',
        '   4.2 Dependency Injection (GetIt)',
        '   4.3 BLoC State Management Pattern',
        '   4.4 Error Handling (Exception/Failure/Either)',
        '   4.5 Repository Pattern',
        '   4.6 Use Case Pattern',
        '   4.7 Multi-Layer Caching Strategy',
        '   4.8 Feature Flags (Compile-Time + Runtime)',
        '   4.9 Navigation & Routing',
        '   4.10 Feature Modules Reference',
        '5. Backend Architecture (Cloud Functions)',
        '   5.1 Entry Point & Module Organization',
        '   5.2 Authentication & Permission Model',
        '   5.3 Error Handling & Response Standardization',
        '   5.4 Firestore Transactions & Batch Operations',
        '   5.5 Scheduled Functions & Cron Jobs',
        '   5.6 Firestore Triggers',
        '   5.7 Retry & Idempotency Patterns',
        '   5.8 Concurrency & Race Condition Prevention',
        '   5.9 Validation Patterns',
        '   5.10 Logging & Observability',
        '   5.11 Function Catalog (All 200+)',
        '6. Database Schema (Firestore)',
        '   6.1 Core Collections & Document Structures',
        '   6.2 Subcollections',
        '   6.3 Composite Indexes (85+)',
        '   6.4 Field Overrides',
        '7. Feature Engineering Deep Dives',
        '   7.1 Discovery & Matching Pipeline',
        '   7.2 Chat & Real-Time Messaging',
        '   7.3 Gamification Engine',
        '   7.4 Vocabulary Processing Pipeline',
        '   7.5 Video Calling (WebRTC)',
        '   7.6 Coin Economy & FIFO Expiration',
        '   7.7 Subscription Lifecycle & Webhooks',
        '   7.8 Media Processing Pipeline',
        '   7.9 Safety & AI Moderation',
        '   7.10 Language Learning System',
        '   7.11 Admin Panel & RBAC',
        '   7.12 Notification System (FCM + Brevo)',
        '8. Scalability Engineering',
        '9. Security Architecture',
        '   9.1 Firestore Security Rules (Detailed)',
        '   9.2 Storage Security Rules',
        '   9.3 Firebase App Check',
        '   9.4 Admin RBAC & Audit Trail',
        '   9.5 Security Risks & Recommendations',
        '10. Build & Deployment',
        '   10.1 Android Build Configuration',
        '   10.2 iOS Build Configuration',
        '   10.3 Firebase Deployment',
        '   10.4 Emulator Setup',
        '11. Localization Engineering',
        '12. App Initialization Flow',
        '13. API Reference (Cloud Functions)',
        '14. Diagrams Reference',
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(1)
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════
    doc.add_heading('1. Executive Summary & System Overview', level=1)
    doc.add_paragraph(
        'GreenGo is a production-grade mobile dating and language exchange application. '
        'The system is a full-stack Flutter + Firebase application designed for millions of concurrent users. '
        'It employs Clean Architecture with BLoC pattern on the client, 200+ Cloud Functions (TypeScript/Node.js 22) '
        'on the server, and Cloud Firestore with 170+ collections as the primary database.'
    )

    doc.add_heading('Platform Summary', level=2)
    add_styled_table(doc,
        ['Dimension', 'Detail'],
        [
            ['Frontend', 'Flutter/Dart, Clean Architecture, BLoC + Provider, 41 feature modules'],
            ['Backend', 'Firebase Cloud Functions v2 (Gen2), TypeScript, Node.js 22'],
            ['Database', 'Cloud Firestore (170+ collections, 85+ composite indexes)'],
            ['Auth', 'Firebase Auth (email/password, Google, Apple)'],
            ['Storage', 'Cloud Storage (photos, videos, voice, exports)'],
            ['Notifications', 'FCM (push) + Brevo (email) + SendGrid (legacy)'],
            ['Payments', 'Google Play Billing + Apple StoreKit + Stripe'],
            ['Video', 'WebRTC via Agora + Firestore signaling'],
            ['AI/ML', 'Google Vision, NLP, Cloud Translation, ML Kit (on-device)'],
            ['Monitoring', 'Firebase Analytics, Crashlytics, Performance Monitoring'],
            ['Localization', '7 languages (EN, DE, ES, FR, IT, PT, PT-BR) via ARB'],
            ['CI/CD', 'Firebase CLI deployment, debug/release signing, ProGuard (disabled)'],
            ['Firebase Project', 'greengo-chat'],
            ['Package ID', 'com.greengochat.greengochatapp'],
            ['App Version', '1.1.0+26'],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════
    doc.add_heading('2. System Architecture & Component Diagram', level=1)
    doc.add_paragraph(
        'The system has four architectural layers: Client, Firebase Services, Cloud Functions, and External Services. '
        'See docs/diagrams/system_architecture.drawio for the visual diagram.'
    )

    doc.add_heading('Layer 1: Client (Flutter App)', level=2)
    doc.add_paragraph(
        'The Flutter app communicates with Firebase services via SDKs (Firestore, Auth, Storage, FCM) '
        'and with Cloud Functions via the cloud_functions SDK (onCall). Real-time data flows through '
        'Firestore snapshot listeners. Offline support uses Firestore persistence (unlimited cache) '
        'and Hive for local key-value storage.'
    )

    doc.add_heading('Layer 2: Firebase Services', level=2)
    add_styled_table(doc,
        ['Service', 'SDK Version', 'Purpose', 'Key Configuration'],
        [
            ['Cloud Firestore', '5.0.0', 'Primary database', 'Offline persistence enabled, unlimited cache size'],
            ['Firebase Auth', '5.0.0', 'Authentication', 'Email/password (Google/Apple disabled for MVP)'],
            ['Cloud Storage', '12.0.0', 'File storage', 'Open rules (dev) - MUST lock down for production'],
            ['Cloud Functions', '5.0.0', 'Server logic', '200+ functions, v2 Gen2 + legacy v1'],
            ['FCM', '15.0.0', 'Push notifications', 'Background + foreground handling'],
            ['Analytics', '11.0.0', 'Event tracking', 'Custom events + user properties'],
            ['Crashlytics', '4.0.0', 'Crash reporting', 'Disabled in emulator mode'],
            ['Performance', '0.10.0', 'Perf monitoring', 'Disabled in emulator mode'],
            ['Remote Config', '5.0.0', 'Feature flags', '10s fetch (debug), 1hr (prod)'],
            ['App Check', '0.3.0', 'API protection', 'Play Integrity (Android), App Attest (iOS)'],
        ]
    )

    doc.add_heading('Layer 3: Cloud Functions', level=2)
    doc.add_paragraph(
        '200+ functions organized into 16 service modules. Entry point: functions/src/index.ts. '
        'Local development uses index-minimal.ts (subset). CRITICAL: Never deploy with --force.'
    )

    doc.add_heading('Layer 4: External Services', level=2)
    add_styled_table(doc,
        ['Service', 'Package', 'Purpose'],
        [
            ['Brevo (Sendinblue)', 'brevo SDK', 'Transactional email, digests, re-engagement'],
            ['SendGrid', '@sendgrid/mail 8.1.0', 'Legacy email (being replaced by Brevo)'],
            ['Google Play Billing', 'googleapis', 'Android IAP verification'],
            ['Apple StoreKit', 'app-store-server-library 2.0.0', 'iOS IAP verification'],
            ['Stripe', 'stripe 14.8.0', 'Payment processing'],
            ['Agora', 'agora-access-token 2.0.4', 'WebRTC video call tokens'],
            ['Twilio', 'twilio 4.19.3', 'SMS/voice (2FA, etc.)'],
            ['Google Cloud Vision', '@google-cloud/vision', 'Image moderation (NSFW)'],
            ['Google Cloud NLP', '@google-cloud/language', 'Text moderation'],
            ['Google Cloud Translation', '@google-cloud/translate', 'Auto-translation'],
            ['Google Cloud Speech', '@google-cloud/speech', 'Voice transcription'],
            ['Sharp', 'sharp 0.33.5', 'Image compression & thumbnails'],
            ['FFmpeg', 'fluent-ffmpeg 2.1.2', 'Video processing'],
            ['PDFKit', 'pdfkit 0.14.0', 'PDF export generation'],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 3. TECHNOLOGY STACK
    # ══════════════════════════════════════════════════
    doc.add_heading('3. Technology Stack & Dependencies', level=1)

    doc.add_heading('3.1 Frontend Dependencies (pubspec.yaml)', level=2)
    add_styled_table(doc,
        ['Category', 'Package', 'Version', 'Purpose'],
        [
            ['State Mgmt', 'flutter_bloc', '8.1.3', 'BLoC pattern (events + states)'],
            ['State Mgmt', 'provider', '6.1.1', 'Language switching, DI wrapper'],
            ['State Mgmt', 'equatable', '2.0.5', 'Value equality for BLoC states'],
            ['DI', 'get_it', '7.6.4', 'Service locator (singleton)'],
            ['DI', 'injectable', '2.3.2', 'Code generation for get_it registrations'],
            ['FP', 'dartz', '0.10.1', 'Either<L,R> for functional error handling'],
            ['Validation', 'formz', '0.6.1', 'Form input validation'],
            ['Network', 'dio', '5.4.0', 'HTTP client with interceptors'],
            ['Storage', 'hive / hive_flutter', '2.2.3', 'Local encrypted key-value store'],
            ['Storage', 'shared_preferences', '2.2.2', 'Simple key-value persistence'],
            ['Media', 'image_picker', '1.0.4', 'Camera/gallery photo selection'],
            ['Media', 'flutter_image_compress', '2.1.0', 'Client-side image compression'],
            ['Media', 'video_player', '2.8.2', 'Video playback'],
            ['Media', 'record', '5.2.1', 'Audio recording'],
            ['Media', 'audioplayers', '6.1.0', 'Audio playback'],
            ['Media', 'flutter_tts', '4.2.0', 'Text-to-speech (pronunciation)'],
            ['ML Kit', 'google_mlkit_face_detection', '0.13.1', 'On-device face detection'],
            ['ML Kit', 'google_mlkit_text_recognition', '0.15.0', 'On-device OCR'],
            ['ML Kit', 'google_mlkit_image_labeling', '0.14.0', 'On-device image classification'],
            ['IAP', 'in_app_purchase', '3.2.0', 'Cross-platform purchase API'],
            ['Maps', 'google_maps_flutter', '2.5.0', 'Map display'],
            ['Maps', 'geolocator', '10.1.0', 'GPS location'],
            ['UI', 'lottie', '2.7.0', 'Lottie animations (celebrations)'],
            ['UI', 'shimmer', '3.0.0', 'Loading skeleton effects'],
            ['UI', 'cached_network_image', '3.3.0', 'Cached image loading'],
        ]
    )

    doc.add_heading('3.2 Backend Dependencies (functions/package.json)', level=2)
    add_styled_table(doc,
        ['Category', 'Package', 'Version', 'Purpose'],
        [
            ['Firebase', 'firebase-admin', '13.6.1', 'Admin SDK (Firestore, Auth, Storage)'],
            ['Firebase', 'firebase-functions', '7.0.5', 'Cloud Functions v2 framework'],
            ['Validation', 'joi', '17.11.0', 'Schema validation'],
            ['Validation', 'zod', '3.22.4', 'TypeScript-first schema validation'],
            ['Validation', 'validator', '13.11.0', 'String validators (email, URL)'],
            ['Auth', 'bcryptjs', '2.4.3', 'Password hashing'],
            ['Auth', 'jsonwebtoken', '9.0.2', 'JWT token generation/verification'],
            ['HTTP', 'express', '4.18.2', 'HTTP server (webhook endpoints)'],
            ['HTTP', 'axios', '1.6.2', 'HTTP client'],
            ['HTTP', 'cors', '2.8.5', 'Cross-origin resource sharing'],
            ['Image', 'sharp', '0.33.5', 'Image compression, resize, thumbnails'],
            ['Video', 'fluent-ffmpeg', '2.1.2', 'Video transcoding'],
            ['PDF', 'pdfkit', '0.14.0', 'PDF generation for exports'],
            ['Util', 'lodash', '4.17.21', 'Array/object utilities, chunk()'],
            ['Util', 'uuid', '9.0.1', 'UUID generation'],
            ['Util', 'date-fns', '3.0.1', 'Date manipulation'],
            ['TypeScript', 'typescript', '5.3.3', 'TypeScript compiler'],
            ['Lint', 'eslint (google)', '-', 'Code style enforcement'],
            ['Test', 'jest + ts-jest', '-', 'Unit testing framework'],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 4. FRONTEND ARCHITECTURE
    # ══════════════════════════════════════════════════
    doc.add_heading('4. Frontend Architecture', level=1)

    # 4.1
    doc.add_heading('4.1 Clean Architecture Pattern', level=2)
    doc.add_paragraph(
        'Each feature module follows a strict three-layer architecture with unidirectional dependencies: '
        'Presentation -> Domain -> Data. The domain layer has ZERO framework dependencies.'
    )
    add_code_block(doc, '''lib/features/{feature}/
  domain/
    entities/        # Pure Dart classes (no Firebase imports)
    repositories/    # Abstract interfaces (contracts)
    usecases/        # Business logic (single responsibility)
  data/
    models/          # Firestore serialization (fromFirestore/toFirestore)
    datasources/     # Remote (Firestore) + Local (Hive) data sources
    repositories/    # Concrete implementations of domain interfaces
  presentation/
    bloc/            # BLoC (events, states, bloc class)
    screens/         # Full-page widgets
    widgets/         # Reusable UI components''')

    # 4.2
    doc.add_heading('4.2 Dependency Injection (GetIt)', level=2)
    doc.add_paragraph(
        'File: lib/core/di/injection_container.dart. Uses GetIt service locator with three registration strategies:'
    )
    add_code_block(doc, '''final sl = GetIt.instance;  // Global service locator

// Registration strategies:
sl.registerFactory(() => ChatBloc(sl(), sl(), sl()));     // New instance per request (BLoCs)
sl.registerLazySingleton(() => SendMessage(sl()));         // Singleton, lazy-loaded (Use Cases)
sl.registerLazySingleton<ChatRepository>(                  // Interface -> Implementation
  () => ChatRepositoryImpl(remoteDataSource: sl()),
);

// Conditional registration based on compile-time feature flags:
if (AppConfig.enableVideoCalls) {
  sl.registerFactory(() => VideoCallBloc(sl()));
  sl.registerLazySingleton<VideoCallingRepository>(
    () => VideoCallingRepositoryImpl(remoteDataSource: sl()),
  );
}''')
    doc.add_paragraph(
        'Registration order: External services (Firebase, SharedPrefs) -> Data Sources -> '
        'Repositories -> Use Cases -> BLoCs. All dependencies are lazy-loaded singletons except '
        'BLoCs which are factories (new instance per widget).'
    )

    # 4.3
    doc.add_heading('4.3 BLoC State Management Pattern', level=2)
    doc.add_paragraph(
        'Each feature uses flutter_bloc with typed events and states. Events are immutable input commands; '
        'states are immutable output snapshots with copyWith() for immutable updates.'
    )
    doc.add_heading('Event Pattern', level=3)
    add_code_block(doc, '''abstract class ChatEvent { const ChatEvent(); }

class ChatConversationLoaded extends ChatEvent {
  final String matchId;
  final String currentUserId;
  final String otherUserId;
  const ChatConversationLoaded({required this.matchId, ...});
}

class ChatMessageSent extends ChatEvent {
  final String content;
  final MessageType type;
  final MembershipRules? membershipRules;  // For usage limit checks
  const ChatMessageSent({required this.content, ...});
}''')

    doc.add_heading('State Pattern', level=3)
    add_code_block(doc, '''abstract class ChatState { const ChatState(); }

class ChatInitial extends ChatState {}
class ChatLoading extends ChatState {}
class ChatLoaded extends ChatState {
  final Conversation conversation;
  final List<Message> messages;
  final bool isOtherUserTyping;

  // Immutable update via copyWith:
  ChatLoaded copyWith({List<Message>? messages, bool? isOtherUserTyping}) {
    return ChatLoaded(
      conversation: conversation,
      messages: messages ?? this.messages,
      isOtherUserTyping: isOtherUserTyping ?? this.isOtherUserTyping,
    );
  }
}
class ChatError extends ChatState { final String message; }
class ChatMessageLimitReached extends ChatState { final UsageLimitResult limitResult; }''')

    doc.add_heading('BLoC Handler Pattern', level=3)
    add_code_block(doc, '''class ChatBloc extends Bloc<ChatEvent, ChatState> {
  ChatBloc({required this.sendMessage, ...}) : super(const ChatInitial()) {
    on<ChatConversationLoaded>(_onConversationLoaded);
    on<ChatMessageSent>(_onMessageSent);
  }

  Future<void> _onConversationLoaded(event, Emitter<ChatState> emit) async {
    emit(const ChatLoading());
    // Stream real-time Firestore messages with emit.forEach:
    await emit.forEach(
      getMessages(GetMessagesParams(conversationId: _conversationId)),
      onData: (messagesResult) => messagesResult.fold(
        (failure) => ChatError(failure.toString()),
        (messages) => ChatLoaded(conversation: conv, messages: messages),
      ),
    );
  }

  Future<void> _onMessageSent(event, Emitter<ChatState> emit) async {
    // 1. Check membership usage limits
    final limit = await _usageLimitService.checkLimit(...);
    if (!limit.isAllowed) { emit(ChatMessageLimitReached(...)); return; }
    // 2. Execute use case
    final result = await sendMessage(SendMessageParams(...));
    result.fold(
      (failure) => emit(ChatError(failure.toString())),
      (_) => _resubscribeToMessages(),
    );
  }
}''')

    # 4.4
    doc.add_heading('4.4 Error Handling (Exception/Failure/Either)', level=2)
    doc.add_paragraph(
        'Three-layer error model: Exceptions (data layer) -> Failures (domain layer) -> Either<Failure,T> (use cases).'
    )
    doc.add_heading('Exceptions (Data Layer)', level=3)
    add_code_block(doc, '''class ServerException implements Exception { final String message; }
class CacheException implements Exception { final String message; }
class NetworkException implements Exception { final String message; }
class AuthenticationException implements Exception { final String message; }
class UploadException implements Exception { final String message; }
class ValidationException implements Exception { final String message; }''')

    doc.add_heading('Failures (Domain Layer - Equatable)', level=3)
    add_code_block(doc, '''abstract class Failure extends Equatable {
  final String message;
  const Failure(this.message);
  @override List<Object?> get props => [message];
}

class ServerFailure extends Failure { const ServerFailure(super.message); }
class AuthenticationFailure extends Failure { ... }
class InvalidCredentialsFailure extends Failure { ... }
class EmailAlreadyInUseFailure extends Failure { ... }
// ... domain-specific failures''')

    doc.add_heading('Repository Exception-to-Failure Conversion', level=3)
    add_code_block(doc, '''// In repository implementation:
Future<Either<Failure, User>> signIn(String email, String password) async {
  try {
    final user = await remoteDataSource.signInWithEmail(email, password);
    return Right(user);           // Success path
  } on AuthenticationException catch (e) {
    return Left(AuthenticationFailure(e.message));  // Known error
  } catch (e) {
    return Left(ServerFailure(e.toString()));       // Unknown error
  }
}

// In BLoC handler:
final result = await signIn(params);
result.fold(
  (failure) => emit(AuthError(failure.message)),  // Left: error
  (user) => emit(AuthSuccess(user)),              // Right: success
);''')

    # 4.5
    doc.add_heading('4.5 Repository Pattern', level=2)
    add_code_block(doc, '''// Domain layer (interface):
abstract class DiscoveryRepository {
  Future<Either<Failure, List<MatchCandidate>>> getDiscoveryStack({
    required String userId,
    required MatchPreferences preferences,
    int limit = 20,
    bool forceRefresh = false,
  });
}

// Data layer (implementation):
class DiscoveryRepositoryImpl implements DiscoveryRepository {
  final DiscoveryRemoteDataSource remoteDataSource;
  static const int queueSize = 20;

  @override
  Future<Either<Failure, List<MatchCandidate>>> getDiscoveryStack({...}) async {
    try {
      final candidates = await remoteDataSource.getDiscoveryStack(...);
      return Right(candidates);
    } on ServerException catch (e) {
      return Left(ServerFailure(e.message));
    } catch (e) {
      if (e.toString().contains('User profile not found')) return const Right([]);
      return Left(ServerFailure(e.toString()));
    }
  }
}''')

    # 4.6
    doc.add_heading('4.6 Use Case Pattern', level=2)
    add_code_block(doc, '''// Base class:
abstract class UseCase<Type, Params> {
  Future<Either<Failure, Type>> call(Params params);
}
class NoParams { const NoParams(); }

// Concrete use case:
class SendMessage {
  final ChatRepository repository;
  SendMessage(this.repository);

  Future<Either<Failure, Message>> call(SendMessageParams params) {
    return repository.sendMessage(
      matchId: params.matchId, senderId: params.senderId,
      receiverId: params.receiverId, content: params.content,
      type: params.type, metadata: params.metadata,
    );
  }
}''')

    # 4.7
    doc.add_heading('4.7 Multi-Layer Caching Strategy', level=2)
    doc.add_paragraph(
        'File: lib/core/services/cache_service.dart. Implements two-tier caching: '
        'in-memory (fast, max 100 entries, LRU eviction) + Hive persistent storage. '
        'Each data type has a specific TTL.'
    )
    add_styled_table(doc,
        ['Data Type', 'TTL', 'Storage', 'Eviction'],
        [
            ['Discovery Stack', '3 minutes', 'Memory + Hive', 'TTL-based'],
            ['User Profiles', '10 minutes', 'Memory + Hive (profiles_cache box)', 'TTL-based'],
            ['Matches', '5 minutes', 'Memory + Hive', 'TTL-based'],
            ['Settings', '1 hour', 'Memory + Hive (settings_cache box)', 'TTL-based'],
            ['Admin Candidate', '1 hour', 'Memory only (session)', 'Session-scoped'],
            ['General Data', '5 minutes (default)', 'Memory + Hive (general_cache box)', 'TTL + LRU'],
        ]
    )
    add_code_block(doc, '''// Two-tier lookup pattern:
T? get<T>(String key) {
  // 1. Check in-memory cache first (O(1) HashMap lookup)
  final memEntry = _memoryCache[key];
  if (memEntry != null && !memEntry.isExpired) return memEntry.data as T?;

  // 2. Fall back to persistent Hive storage
  final cached = _generalBox.get(key);
  if (cached != null) {
    final wrapper = jsonDecode(cached);
    final cachedAt = DateTime.fromMillisecondsSinceEpoch(wrapper['_cachedAt']);
    if (DateTime.now().difference(cachedAt) < ttl) {
      _setMemoryCache(key, wrapper['data'], ttl);  // Promote to memory
      return wrapper['data'] as T?;
    }
  }
  return null;
}

// LRU eviction when memory cache exceeds 100 entries:
void _setMemoryCache(String key, dynamic data, Duration ttl) {
  if (_memoryCache.length >= maxMemoryCacheSize) {
    final oldest = _memoryCache.keys.toList()
      ..sort((a, b) => _memoryCache[a]!.createdAt.compareTo(_memoryCache[b]!.createdAt));
    for (int i = 0; i < 10; i++) _memoryCache.remove(oldest[i]);
  }
  _memoryCache[key] = _CacheEntry(data, ttl);
}''')

    # 4.8
    doc.add_heading('4.8 Feature Flags (Compile-Time + Runtime)', level=2)

    doc.add_heading('Compile-Time Flags (AppConfig)', level=3)
    add_code_block(doc, '''class AppConfig {
  AppConfig._();
  static String get environment => kDebugMode ? 'Development' : 'Production';

  // Compile-time feature gates:
  static const bool enableVideoCalls = false;       // NDK issues
  static const bool enableVoiceMessages = false;
  static const bool enableLanguageLearning = true;
  static const bool enableGamification = true;
  static const bool enableInAppPurchases = true;
  static const bool enableGoogleAuth = false;       // MVP disabled
  static const bool enableAppleAuth = false;        // MVP disabled

  // Emulator configuration:
  static const bool _forceEmulators = bool.fromEnvironment('USE_EMULATORS');
  static bool get useLocalEmulators => kDebugMode && _forceEmulators;
  static const String emulatorHost = '10.0.2.2';   // Android emulator
}''')

    doc.add_heading('Runtime Flags (Firestore-backed)', level=3)
    add_code_block(doc, '''class FeatureFlagsService extends ChangeNotifier {
  // Singleton pattern
  static final _instance = FeatureFlagsService._internal();
  factory FeatureFlagsService() => _instance;

  // Real-time Firestore listener on app_config/feature_flags
  StreamSubscription<DocumentSnapshot>? _subscription;
  Map<String, bool> _flags = {};

  // Hardcoded fallback defaults (20+ flags):
  static const Map<String, bool> _defaults = {
    'discovery': true, 'matches': true, 'messaging': true,
    'videoCalls': false, 'coins': true, 'gamification': true, ...
  };

  Future<void> initialize() async {
    final docRef = _firestore.doc('app_config/feature_flags');
    final docSnap = await docRef.get();
    if (docSnap.exists) _updateFlags(docSnap);
    else _flags = Map.from(_defaults);

    // Real-time listener for admin toggle changes:
    _subscription = docRef.snapshots().listen((snapshot) {
      _updateFlags(snapshot);
      notifyListeners();  // Triggers UI rebuild
    });
  }

  bool isEnabled(String feature) => _flags[feature] ?? _defaults[feature] ?? false;
}''')

    # 4.9
    doc.add_heading('4.9 Navigation & Routing', level=2)
    doc.add_paragraph(
        'Uses native Flutter Navigator with named routes defined in main.dart. '
        'MainNavigationScreen provides bottom tab navigation. safe_navigation.dart '
        'wraps Navigator calls with error handling. No GoRouter or GetX.'
    )

    # 4.10
    doc.add_heading('4.10 Feature Modules Reference (41 Total)', level=2)
    add_styled_table(doc,
        ['Module', 'Architecture', 'Description'],
        [
            ['authentication', 'Domain+Data+Presentation', 'Login, registration, password reset, 2FA'],
            ['discovery', 'Domain+Data+Presentation', 'Swipe-based matching with scoring algorithm'],
            ['chat', 'Domain+Data+Presentation', 'Real-time messaging (text/image/voice/video/gif)'],
            ['gamification', 'Domain+Data+Presentation', 'XP, levels, achievements, challenges, vocab'],
            ['profile', 'Domain+Data+Presentation', 'Profile CRUD, photo management, verification'],
            ['video_calling', 'Domain+Data+Presentation', 'WebRTC 1:1 and group calls'],
            ['coins', 'Domain+Data+Presentation', 'Virtual currency, purchases, gifting'],
            ['membership', 'Domain+Data+Presentation', 'Subscription tiers, IAP flow'],
            ['notifications', 'Domain+Data+Presentation', 'FCM handling, preference management'],
            ['language_learning', 'Domain+Data+Presentation', 'Lessons, flashcards, constellation'],
            ['admin', 'Domain+Data+Presentation', 'Admin panel integration'],
            ['safety', 'Domain+Presentation', 'Safety tips, reporting UI'],
            ['matching', 'Domain only', 'Match scoring entities and algorithm'],
        ]
    )
    doc.add_paragraph('Plus 28 more feature modules (blind_date, communities, events, etc.)')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 5. BACKEND ARCHITECTURE
    # ══════════════════════════════════════════════════
    doc.add_heading('5. Backend Architecture (Cloud Functions)', level=1)

    # 5.1
    doc.add_heading('5.1 Entry Point & Module Organization', level=2)
    add_code_block(doc, '''// functions/src/index.ts - Main entry point
// Re-exports all functions from service modules:
export { compressUploadedImage, compressImage } from './media/imageCompression';
export { translateMessage, autoTranslateMessage, ... } from './messaging/translation';
export { grantXP, trackAchievementProgress, ... } from './gamification/gamificationManager';
export { onMessageCreatedVocabulary, computeDailyUserStats, refreshMyStats } from './gamification/index';
// ... 30+ export blocks covering 200+ functions

// functions/src/index-minimal.ts - Local development subset
// CRITICAL: Never deploy with --force (deletes 100+ production functions not in local source)''')

    doc.add_heading('Directory Structure', level=3)
    add_code_block(doc, '''functions/src/
  shared/
    firebaseAdmin.ts     # Firebase Admin SDK initialization (singleton)
    utils.ts             # Auth verification, error handling, logging, validation, retry
    types.ts             # TypeScript interfaces and enums
    purchase_verification.ts  # Google Play + App Store receipt verification
  admin/
    adminDashboard.ts    # Dashboard metrics (DAU, revenue, etc.)
    roleManagement.ts    # RBAC with 4 roles, 30+ permissions
    userManagement.ts    # User CRUD, suspension, banning
    moderationQueue.ts   # Content moderation queue
    adminPanelFunctions.ts  # 2FA, password mgmt, AI support
    mvp_access.ts        # Early access user approval
  gamification/
    index.ts             # v2 functions (XP, achievements, vocabulary, stats)
    gamificationManager.ts  # Legacy v1 functions (same 8 endpoints)
    handlers.ts          # Pure business logic (testable without Firebase)
    userStatsCompute.ts  # Batch stats computation with cursor pagination
    vocabularyProcessor.ts  # Firestore trigger for word extraction
  media/                 # Image compression, video processing, voice transcription
  messaging/             # Translation, scheduled messages
  safety/                # Photo/text moderation, reporting, identity verification
  analytics/             # Revenue, cohort, churn prediction, A/B testing
  notifications/         # FCM push + Brevo email + triggers
  video_calling/         # 1:1 + group calls + features
  coins/                 # Virtual currency management
  subscription/          # Membership verification + expiration
  subscriptions/         # Webhook handlers (Play Store + App Store)
  language_learning/     # Lesson CRUD, teacher management, progress
  discovery/             # Candidate pool precomputation
  presence/              # Online status tracking + cleanup
  security/              # Security audit
  backup/                # Conversation backup + PDF export''')

    # 5.2
    doc.add_heading('5.2 Authentication & Permission Model', level=2)
    add_code_block(doc, '''// shared/utils.ts - Two-level auth:

// Level 1: Basic user authentication
export async function verifyAuth(auth: AuthData | undefined): Promise<string> {
  if (!auth?.uid) {
    throw new HttpsError('unauthenticated', 'User must be authenticated');
  }
  return auth.uid;
}

// Level 2: Admin authentication (checks Firestore admin_users collection)
export async function verifyAdminAuth(context: CallableContext): Promise<string> {
  const uid = await verifyAuth(context.auth);
  const adminDoc = await db.collection('admin_users').doc(uid).get();
  if (!adminDoc.exists) {
    throw new HttpsError('permission-denied', 'Admin access required');
  }
  return uid;
}

// Level 3: Permission-specific check (in admin functions)
async function verifyAdminPermission(context, requiredPermission: string) {
  const uid = await verifyAdminAuth(context);
  const adminDoc = await db.collection('admin_users').doc(uid).get();
  const permissions = adminDoc.data()?.permissions || [];
  if (!permissions.includes(requiredPermission)) {
    throw new HttpsError('permission-denied', `Missing permission: ${requiredPermission}`);
  }
  // Log to audit trail:
  await db.collection('admin_audit_log').add({
    adminId: uid, action: requiredPermission,
    timestamp: FieldValue.serverTimestamp(), ipAddress: context.rawRequest?.ip
  });
}''')

    doc.add_heading('RBAC Roles & Permissions', level=3)
    add_styled_table(doc,
        ['Role', 'Permissions'],
        [
            ['superAdmin', 'ALL 30+ permissions including manageAdmins, systemSettings, deleteUsers'],
            ['moderator', 'viewReports, reviewContent, issueWarnings, suspendUsers, banUsers'],
            ['support', 'viewUserProfiles, editUserProfiles, viewSubscriptions, adjustCoins'],
            ['analyst', 'viewAnalytics, exportData, viewAuditLog'],
        ]
    )

    # 5.3
    doc.add_heading('5.3 Error Handling & Response Standardization', level=2)
    add_code_block(doc, '''// Custom error class with HTTP status mapping:
class AppError extends Error {
  constructor(
    public code: string,
    public message: string,
    public statusCode: number = 500,
    public details?: any
  ) { super(message); }
}

// Converts AppError to Firebase HttpsError:
export function handleError(error: unknown): HttpsError {
  if (error instanceof AppError) {
    const code = getHttpsErrorCode(error.statusCode);
    return new HttpsError(code, error.message, error.details);
  }
  return new HttpsError('internal', 'An unexpected error occurred');
}

// HTTP status -> Firebase error code mapping:
function getHttpsErrorCode(status: number): FunctionsErrorCode {
  switch (status) {
    case 400: return 'invalid-argument';
    case 401: return 'unauthenticated';
    case 403: return 'permission-denied';
    case 404: return 'not-found';
    case 409: return 'already-exists';
    case 429: return 'resource-exhausted';
    case 503: return 'unavailable';
    default: return 'internal';
  }
}

// Standard API response pattern:
interface ApiResponse<T = any> {
  success: boolean;
  data?: T;
  error?: { code: string; message: string; details?: any; };
  message?: string;
}''')

    # 5.4
    doc.add_heading('5.4 Firestore Transactions & Batch Operations', level=2)

    doc.add_heading('Transactions (Multi-Document Consistency)', level=3)
    add_code_block(doc, '''// Used for: coin operations, subscription state changes, XP grants
// Example: Coin purchase with batch expiry tracking
await db.runTransaction(async (transaction) => {
  // 1. Read current state WITHIN transaction (consistent snapshot)
  const balanceRef = db.collection('coinBalances').doc(userId);
  const balanceDoc = await transaction.get(balanceRef);
  let currentBalance = balanceDoc.data()?.totalCoins || 0;
  let coinBatches = balanceDoc.data()?.coinBatches || [];

  // 2. Calculate changes
  coinBatches.push({
    batchId: uuid(), initialCoins: amount, remainingCoins: amount,
    source: 'purchase', expirationDate: Timestamp.fromDate(expiryDate),
  });

  // 3. Atomic multi-doc update (all succeed or all fail)
  transaction.set(balanceRef, { userId, totalCoins: currentBalance + amount, coinBatches });
  transaction.set(db.collection('coinTransactions').doc(), {
    userId, type: 'credit', amount, createdAt: FieldValue.serverTimestamp()
  });
});''')

    doc.add_heading('Batch Writes (Bulk Updates)', level=3)
    add_code_block(doc, '''// Used for: pool precomputation, presence cleanup, vocabulary processing
// Rule: Max 500 write operations per batch.commit()
const BATCH_LIMIT = 450;  // Leave headroom

for (let i = 0; i < items.length; i += BATCH_LIMIT) {
  const batch = db.batch();
  const chunk = items.slice(i, i + BATCH_LIMIT);
  for (const item of chunk) {
    const ref = db.collection('candidate_pools').doc(item.key);
    batch.set(ref, { members: item.members, count: item.count, updatedAt: now });
  }
  await batch.commit();
}''')

    # 5.5
    doc.add_heading('5.5 Scheduled Functions & Cron Jobs', level=2)
    add_styled_table(doc,
        ['Function', 'Schedule (UTC)', 'Memory', 'Timeout', 'Purpose'],
        [
            ['resetDailyChallenges', '0 0 * * * (midnight)', '512MiB', '300s', 'Reset daily challenge progress'],
            ['updateLeaderboardRankings', '0 * * * * (hourly)', '512MiB', '300s', 'Recompute XP leaderboards'],
            ['computeDailyUserStats', '0 3 * * * (3 AM)', '1GiB', '540s', 'Batch user stats (cursor pagination)'],
            ['cleanupDisappearingMedia', 'Hourly', '256MiB', '300s', 'Delete expired media'],
            ['sendScheduledMessages', 'Every minute', '256MiB', '60s', 'Send pending scheduled messages'],
            ['processExpiredCoins', '0 2 * * * (2 AM)', '256MiB', '300s', 'FIFO coin batch expiration'],
            ['sendExpirationWarnings', '0 10 * * * (10 AM)', '256MiB', '300s', 'Warn users of expiring coins'],
            ['grantMonthlyAllowances', '0 0 1 * * (1st)', '512MiB', '300s', 'Monthly tier coin grants'],
            ['scheduledChurnPrediction', '0 2 * * * (2 AM)', '1GiB', '540s', 'ML churn prediction batch'],
            ['cleanupStalePresence', 'Every 5 min', '256MiB', '60s', 'Mark stale users offline'],
            ['scheduledSecurityAudit', 'Weekly', '512MiB', '300s', 'Automated security audit'],
            ['candidatePoolPrecompute', 'Every 10 min', '512MiB', '300s', 'Rebuild discovery candidate pools'],
        ]
    )

    # 5.6
    doc.add_heading('5.6 Firestore Triggers', level=2)
    add_styled_table(doc,
        ['Trigger', 'Collection', 'Type', 'Action'],
        [
            ['onMessageCreatedVocabulary', 'conversations/{id}/messages/{msgId}', 'onCreate', 'Extract words, award XP'],
            ['onNewLikePush', 'likes/{likeId}', 'onCreate', 'FCM notification to liked user'],
            ['onNewMatchPush', 'matches/{matchId}', 'onCreate', 'FCM to both users'],
            ['onNewMessagePush', 'conversations/{id}/messages/{msgId}', 'onCreate', 'FCM to receiver'],
            ['autoTranslateMessage', 'conversations/{id}/messages/{msgId}', 'onCreate', 'Auto-translate if needed'],
            ['compressUploadedImage', 'Storage: profile photos', 'onFinalize', 'Compress + generate thumbnail'],
            ['processUploadedVideo', 'Storage: video files', 'onFinalize', 'Transcode + extract thumbnail'],
            ['transcribeVoiceMessage', 'Storage: voice messages', 'onFinalize', 'Speech-to-text'],
            ['onSupportChatCreated', 'support_chats/{id}', 'onCreate', 'Assign to agent, send notification'],
            ['onVerificationStatusChange', 'profiles/{userId}', 'onUpdate', 'Notify user of verification result'],
        ]
    )

    # 5.7
    doc.add_heading('5.7 Retry & Idempotency Patterns', level=2)
    add_code_block(doc, '''// Retry with exponential backoff (shared/utils.ts):
export async function retry<T>(
  fn: () => Promise<T>, maxRetries = 3, delayMs = 1000
): Promise<T> {
  let lastError: any;
  for (let i = 0; i < maxRetries; i++) {
    try { return await fn(); }
    catch (error) {
      lastError = error;
      if (i < maxRetries - 1)
        await new Promise(r => setTimeout(r, delayMs * (i + 1)));  // Linear backoff
    }
  }
  throw lastError;
}

// Idempotency in triggers (imageCompression.ts):
if (filePath.includes('_compressed') || filePath.includes('_thumb')) return;  // Skip re-processing
// Idempotency in moderation: check if already processed before re-analyzing''')

    # 5.8
    doc.add_heading('5.8 Concurrency & Race Condition Prevention', level=2)
    add_styled_table(doc,
        ['Pattern', 'When Used', 'Example'],
        [
            ['Firestore Transactions', 'Multi-doc consistency', 'Coin operations, subscription changes'],
            ['FieldValue.increment()', 'Atomic counters', 'XP totals, coin balances, message counts'],
            ['FieldValue.serverTimestamp()', 'Clock-independent times', 'All createdAt/updatedAt fields'],
            ['Batch writes (450 limit)', 'Bulk updates', 'Pool precompute, presence cleanup, vocab'],
            ['Promise.all() / Promise.allSettled()', 'Parallel queries', 'Discovery (swipes+matches+blocks)'],
            ['Idempotency checks', 'Trigger re-fire safety', 'Media processing (skip if _compressed)'],
        ]
    )

    # 5.9
    doc.add_heading('5.9 Validation Patterns', level=2)
    add_code_block(doc, '''// shared/utils.ts:
validateRequired(params, ['email', 'userId']);   // Throws MISSING_FIELD
validateEmail(email);                             // Regex validation
validateURL(url);                                 // new URL() validation
createPaginationParams(page, pageSize);           // Enforces max 100/page

// Backend also uses:
// - Joi (17.11.0) for complex schema validation
// - Zod (3.22.4) for TypeScript-first schemas
// - validator (13.11.0) for string validation (email, URL, etc.)''')

    # 5.10
    doc.add_heading('5.10 Logging & Observability', level=2)
    add_code_block(doc, '''// shared/utils.ts logging helpers:
logInfo(message, data?)     // console.log with [INFO] prefix
logError(message, error?)   // console.error with [ERROR] prefix
logWarning(message, data?)  // console.warn with [WARN] prefix

// Admin audit trail (all admin actions logged):
await db.collection('admin_audit_log').add({
  adminId, adminEmail, adminRole, action, targetType, targetId,
  details, timestamp: FieldValue.serverTimestamp(), ipAddress
});

// Client-side: debugPrint() instead of print() (stripped in release builds)''')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 6. DATABASE SCHEMA
    # ══════════════════════════════════════════════════
    doc.add_heading('6. Database Schema (Firestore)', level=1)

    doc.add_heading('6.1 Core Collections & Document Structures', level=2)
    doc.add_paragraph(
        '170+ Firestore collections organized by domain. Document IDs use Firebase UID for user-owned '
        'docs and auto-generated IDs for event/action docs.'
    )

    # profiles
    doc.add_heading('profiles {userId}', level=3)
    add_styled_table(doc,
        ['Field', 'Type', 'Required', 'Description'],
        [
            ['displayName', 'String', 'Yes', 'User display name'],
            ['nickname', 'String', 'No', 'Unique nickname (indexed in nicknames collection)'],
            ['dateOfBirth', 'Timestamp', 'Yes', 'Date of birth for age calculation'],
            ['gender', 'String', 'Yes', 'male/female/other'],
            ['sexualOrientation', 'String', 'No', 'Orientation preference'],
            ['accountStatus', 'String', 'Yes', 'active/suspended/deleted'],
            ['photoUrls', 'String[]', 'Yes', 'Profile photo URLs (Cloud Storage)'],
            ['privatePhotoUrls', 'String[]', 'No', 'Album photos (access-controlled)'],
            ['bio', 'String', 'No', 'User biography text'],
            ['interests', 'String[]', 'No', 'Interest tags for matching'],
            ['location', 'Map{lat,lng,city,country}', 'Yes', 'GPS coordinates + geocoded address'],
            ['languages', 'String[]', 'Yes', 'Spoken language codes'],
            ['verificationStatus', 'String', 'Yes', 'pending/approved/rejected'],
            ['membershipTier', 'String', 'Yes', 'basic/silver/gold/platinum'],
            ['membershipEndDate', 'Timestamp', 'No', 'Active membership expiry'],
            ['isOnline', 'Boolean', 'Yes', 'Online status (presence)'],
            ['lastSeen', 'Timestamp', 'No', 'Last seen timestamp'],
            ['isAdmin', 'Boolean', 'Yes', 'Admin flag'],
            ['isBoosted/isIncognito/isTraveler', 'Boolean', 'No', 'Special mode flags with expiry'],
            ['swipeCount', 'Number', 'No', 'Daily swipe counter'],
            ['messagesSent', 'Number', 'No', 'Lifetime message count'],
            ['createdAt/updatedAt', 'Timestamp', 'Yes', 'Timestamps'],
        ]
    )

    # user_levels
    doc.add_heading('user_levels {userId}', level=3)
    add_styled_table(doc,
        ['Field', 'Type', 'Description'],
        [
            ['level', 'Number', 'Current level (DERIVED from totalXP, not authoritative)'],
            ['currentXP', 'Number', 'XP within current level progress bar'],
            ['totalXP', 'Number', 'Lifetime total XP (authoritative)'],
            ['region', 'String', 'Leaderboard region (global or country code)'],
            ['regionalRank', 'Number', 'Rank within region'],
            ['isVIP', 'Boolean', 'VIP status (level >= 10)'],
            ['lastUpdated', 'Timestamp', 'Last XP update timestamp'],
        ]
    )

    # user_stats (materialized view)
    doc.add_heading('user_stats {userId} (Materialized View)', level=3)
    doc.add_paragraph(
        'Cached computed stats. Refreshed daily at 3AM UTC (computeDailyUserStats) '
        'and on-demand via refreshMyStats callable. Avoids expensive real-time aggregation.'
    )
    add_styled_table(doc,
        ['Field', 'Type', 'Source'],
        [
            ['totalXp', 'Number', 'user_levels.totalXP or xp_transactions sum'],
            ['level', 'Number', 'calculateLevel(totalXp) - always derived'],
            ['messagesSent', 'Number', 'profiles.messagesSent'],
            ['totalConversations', 'Number', 'COUNT(conversations where userId1/userId2)'],
            ['wordsPerLanguage', 'Map<string,number>', 'COUNT(user_vocabulary/words) grouped by language'],
            ['wordsLearnedPerLanguage', 'Map<string,number>', 'COUNT(words where useCount >= 3)'],
            ['achievementsUnlocked', 'Number', 'COUNT(user_achievements where isUnlocked)'],
            ['challengesCompleted', 'Number', 'COUNT(user_challenges where isCompleted)'],
            ['dailyActivity', 'Map<string,number>', 'XP transactions per day (last 30 days)'],
            ['updatedAt', 'Timestamp', 'Last refresh timestamp'],
        ]
    )

    # coinBalances
    doc.add_heading('coinBalances {userId}', level=3)
    add_styled_table(doc,
        ['Field', 'Type', 'Description'],
        [
            ['totalCoins', 'Number', 'Current spendable balance'],
            ['earnedCoins', 'Number', 'Earned via achievements/challenges'],
            ['purchasedCoins', 'Number', 'Bought via IAP'],
            ['giftedCoins', 'Number', 'Received as gifts'],
            ['spentCoins', 'Number', 'Total spent'],
            ['coinBatches', 'Object[]', 'FIFO batches: {batchId, initialCoins, remainingCoins, source, acquiredDate, expirationDate}'],
            ['lastUpdated', 'Timestamp', 'Last balance change'],
        ]
    )

    doc.add_heading('6.2 Subcollections', level=2)
    add_styled_table(doc,
        ['Parent', 'Subcollection', 'Doc ID', 'Key Fields'],
        [
            ['conversations/{id}', 'messages', 'autoId', 'senderId, content, type, sentAt, readAt'],
            ['user_vocabulary/{uid}', 'words', '{lang}_{word}', 'word, language, useCount, firstUsedAt'],
            ['usageLimits/{uid}', 'hours / days', 'YYYY-MM-DD-HH', 'likes, nopes, superLikes, timestamp'],
            ['game_rooms/{id}', 'rounds', 'autoId', 'roundNumber, currentQuestion, answers'],
            ['game_rooms/{id}', 'chat', 'autoId', 'senderId, message, sentAt'],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 7. FEATURE ENGINEERING DEEP DIVES
    # ══════════════════════════════════════════════════
    doc.add_heading('7. Feature Engineering Deep Dives', level=1)

    # 7.1 Discovery
    doc.add_heading('7.1 Discovery & Matching Pipeline', level=2)
    doc.add_paragraph('File: lib/features/discovery/data/datasources/discovery_remote_datasource.dart')
    doc.add_heading('Pipeline Steps', level=3)
    numbered(doc, 'Check session cache (skip to step 7 if valid)')
    numbered(doc, 'Parallel Firestore queries via Future.wait():')
    bullet(doc, 'getSwipeHistory: swipes WHERE userId = X AND timestamp >= 90 days ago ORDER BY timestamp DESC LIMIT 2000')
    bullet(doc, 'getMatchedUserIds: matches WHERE userId1 = X AND isActive = true LIMIT 1000 (+ userId2 direction)')
    bullet(doc, 'getBlockedUserIds: blocked_users WHERE blockerId = X')
    numbered(doc, 'Build exclusion set = swiped + matched + blocked user IDs')
    numbered(doc, 'Query eligible profiles matching user preferences (age, gender, distance, accountStatus=active, isComplete=true)')
    numbered(doc, 'Compute match scores per candidate')
    numbered(doc, 'Sort by score DESC, limit to 20')
    numbered(doc, 'Cache results (session-scoped)')

    doc.add_heading('Match Scoring Algorithm', level=3)
    add_styled_table(doc,
        ['Factor', 'Weight', 'Calculation'],
        [
            ['Location', '40%', 'Inverse distance (closer = higher). Uses Haversine formula.'],
            ['Age Compatibility', '20%', 'Score = 1 - (ageDiff / preferenceRange). Penalizes outside pref.'],
            ['Interest Overlap', '25%', 'Jaccard coefficient: |intersection| / |union| of interest tags'],
            ['Language Match', '15%', 'Shared language count bonus'],
        ]
    )

    # 7.3 Gamification
    doc.add_heading('7.3 Gamification Engine', level=2)
    doc.add_heading('Level Calculation (CRITICAL)', level=3)
    doc.add_paragraph(
        'Level is ALWAYS derived from totalXP using calculateLevel(). Never stored independently. '
        'Both client (Dart) and server (TypeScript) use identical threshold arrays:'
    )
    add_code_block(doc, '''// LEVEL_XP_REQUIREMENTS = [0, 100, 250, 500, 1000, 2000, 3500, 5500, 8000, 11000, 15000, 20000, 26000, 33000, 41000, 50000]

function calculateLevel(totalXp: number): number {
  let level = 1;
  while (level < LEVEL_XP_REQUIREMENTS.length && totalXp >= LEVEL_XP_REQUIREMENTS[level]) {
    level++;
  }
  return level;
}
// Example: 104 XP -> level 2 (crosses 100 threshold)''')

    doc.add_heading('Handler Architecture (Pure Business Logic)', level=3)
    add_code_block(doc, '''// gamification/handlers.ts - Testable without Firebase:
export async function handleGrantXP({ uid, action, metadata }) {
  // 1. Validate action in XP_ACTIONS map
  const xpAmount = XP_ACTIONS[action];
  if (!xpAmount) throw new AppError('INVALID_ACTION', 'Unknown XP action', 400);

  // 2. Read current state
  const levelDoc = await db.collection('user_levels').doc(uid).get();
  const oldXP = levelDoc.data()?.totalXP || 0;
  const oldLevel = calculateLevel(oldXP);

  // 3. Calculate new state
  const newXP = oldXP + xpAmount;
  const newLevel = calculateLevel(newXP);

  // 4. Update user_levels (merge to preserve other fields)
  await db.collection('user_levels').doc(uid).set({
    totalXP: FieldValue.increment(xpAmount),
    currentXP: FieldValue.increment(xpAmount),
    lastUpdated: FieldValue.serverTimestamp(),
  }, { merge: true });

  // 5. Log XP transaction
  await db.collection('xp_transactions').add({
    userId: uid, actionType: action, xpAmount, createdAt: FieldValue.serverTimestamp()
  });

  // 6. Check level-up -> grant coin rewards
  if (newLevel > oldLevel) {
    const coins = calculateLevelRewards(oldLevel, newLevel);
    if (coins > 0) { /* atomic coin grant via transaction */ }
  }

  return { success: true, xpAwarded: xpAmount, newLevel, totalXP: newXP };
}''')

    # 7.4 Vocabulary
    doc.add_heading('7.4 Vocabulary Processing Pipeline', level=2)
    add_code_block(doc, '''// vocabularyProcessor.ts - Firestore trigger
// Trigger: conversations/{conversationId}/messages/{messageId} onCreate
// Memory: 256MiB | Timeout: 30s

1. Extract message text (content || text field)
2. Get conversation language from parent document
3. Regex extract words: /[a-zA-Z\\u00C0-\\u024F\']+/g
4. Filter: length >= 2, deduplicate, lowercase
5. Batch process in groups of 450 (Firestore 500-op batch limit):
   - For each word, check user_vocabulary/{senderId}/words/{lang}_{word}
   - EXISTS: batch.update({ useCount: FieldValue.increment(1) })
   - NEW: batch.set({ word, language, frequencyScore: 0,
           firstUsedAt: serverTimestamp(), useCount: 1 })
   - Increment newWordCount
6. batch.commit()
7. If newWordCount > 0:
   - Update user_levels: totalXP += newWordCount * 1, currentXP += newWordCount * 1
   - Log xp_transaction: { actionType: 'vocabulary_usage', newWords: count }''')

    # 7.6 Coins
    doc.add_heading('7.6 Coin Economy & FIFO Expiration', level=2)
    doc.add_paragraph(
        'Coins use FIFO (First-In, First-Out) batch expiration. Each purchase/grant creates a "batch" '
        'with an expiration date. When coins are spent, the oldest batch is consumed first. '
        'processExpiredCoins runs daily at 2AM UTC to expire old batches.'
    )

    # 7.7 Subscriptions
    doc.add_heading('7.7 Subscription Lifecycle & Webhooks', level=2)
    add_code_block(doc, '''// subscriptionManager.ts - Webhook flow:
// 1. Google Play / App Store sends base64-encoded notification
// 2. Decode and identify notificationType (1-13)
// 3. Find subscription by purchaseToken
// 4. Route to handler:
//    RENEWED -> update currentPeriodEnd
//    CANCELED -> status = 'canceled'
//    ON_HOLD -> status = 'on_hold'
//    IN_GRACE_PERIOD -> start 7-day grace period
//    EXPIRED -> status = 'expired'
// 5. Send renewal reminder 3 days before expiry

// Purchase verification with graceful degradation:
// If Google/Apple API returns 403/401 (not configured):
//   -> Accept purchase (apiUnavailable: true)
//   -> Log guidance to configure API access
//   -> Distinguishes from genuinely invalid purchases''')

    # 7.8 Media
    doc.add_heading('7.8 Media Processing Pipeline', level=2)
    add_code_block(doc, '''// imageCompression.ts - Storage trigger on file upload:
// 1. Idempotency check: skip if path contains '_compressed' or '_thumb'
// 2. Download to /tmp
// 3. Iterative compression: start quality=80, reduce by 10 until < MAX_FILE_SIZE
//    sharp(file).rotate().resize(1920, 1080, { fit: 'inside' }).jpeg({ quality })
// 4. Generate thumbnail: 200x200 cover crop
//    sharp(file).resize(200, 200, { fit: 'cover', position: 'center' })
// 5. Upload compressed + thumbnail to Storage
// 6. Update Firestore message doc with URLs + dimensions
// 7. Clean up /tmp files (fs.unlinkSync)''')
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 8. SCALABILITY
    # ══════════════════════════════════════════════════
    doc.add_heading('8. Scalability Engineering', level=1)

    add_styled_table(doc,
        ['Technique', 'Implementation', 'Impact'],
        [
            ['Materialized Views', 'user_stats/{uid} refreshed at 3AM + on-demand', 'Avoids 6+ queries per stats page view'],
            ['Cursor Pagination', 'startAfter() with 100 docs/page for batch jobs', 'Handles millions of users without OOM'],
            ['Concurrent Batches', 'Promise.allSettled() with 20 users/batch', 'Maximizes throughput within timeouts'],
            ['Timeout Guards', '8-minute check before 9-min function limit', 'Graceful stop, resume on next run'],
            ['Date-Filtered Queries', 'Swipe history limited to 90 days', 'Bounded query size regardless of user age'],
            ['Parallel Queries', 'Future.wait() for independent Firestore calls', 'Reduces wall-clock time 3x for discovery'],
            ['Session Cache', 'Discovery results cached per session', 'Zero Firestore reads on tab switch'],
            ['Admin Cache', '1-hour TTL for admin candidate profile', 'Saves 1 read per discovery load'],
            ['Query Limits', 'Swipes: 2000, Matches: 1000 (down from 5000/2000)', '60% fewer Firestore reads'],
            ['select() Projections', 'Fetch only needed fields', 'Reduced bandwidth and document reads'],
            ['count() Queries', '.count().get() for aggregations', 'Counts without transferring documents'],
            ['FieldValue.increment()', 'Atomic counters', 'No read-modify-write race conditions'],
            ['Batch Size 450', 'Under Firestore 500-op limit with headroom', 'Safe batch operations at scale'],
            ['Pool Precomputation', 'Candidate pools built every 10 min', 'O(1) discovery instead of O(n) scan'],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 9. SECURITY
    # ══════════════════════════════════════════════════
    doc.add_heading('9. Security Architecture', level=1)

    doc.add_heading('9.1 Firestore Security Rules (Detailed)', level=2)
    doc.add_paragraph('File: firestore.rules (681 lines). Key patterns:')

    add_code_block(doc, '''// Helper functions:
function isSignedIn() { return request.auth != null; }
function isOwner(userId) { return request.auth.uid == userId; }
function isAdmin() {
  return isSignedIn() && get(/databases/$(database)/documents/users/$(request.auth.uid)).data.role == 'admin';
}
function isAdminPanelUser() {
  return isSignedIn() && exists(/databases/$(database)/documents/admin_users/$(request.auth.uid));
}''')

    doc.add_heading('Collection-Level Rules Summary', level=3)
    add_styled_table(doc,
        ['Collection', 'Read', 'Write', 'Notes'],
        [
            ['profiles', 'Signed-in', 'Owner + Admin', 'Core user data'],
            ['conversations', 'Signed-in', 'Signed-in', 'No participant check (loose)'],
            ['matches', 'Signed-in', 'Signed-in', 'No participant check (loose)'],
            ['subscriptions', 'Signed-in', 'Server-only', 'Admin SDK writes only'],
            ['user_vocabulary/{uid}/**', 'Owner', 'Owner (create/update)', 'Vocabulary is private'],
            ['admin_users', 'Admin panel', 'Admin panel', 'Restricted to admin role'],
            ['coinBalances', 'Owner', 'Owner + Server', 'Balance can be read by owner'],
            ['Catch-all {document=**}', 'Signed-in', 'Admin panel only', 'Default for unlisted collections'],
        ]
    )

    doc.add_heading('9.2 Storage Security Rules', level=2)
    doc.add_paragraph(
        'CRITICAL: Storage rules are currently FULLY OPEN (allow read, write: if true). '
        'This is for development only. Production deployment MUST add authentication checks.'
    )

    doc.add_heading('9.5 Security Risks & Recommendations', level=2)
    add_styled_table(doc,
        ['Risk', 'Severity', 'Recommendation'],
        [
            ['Storage rules fully open', 'CRITICAL', 'Add auth checks: allow read,write: if request.auth != null'],
            ['matches/conversations loose rules', 'HIGH', 'Add participant check: userId1 == uid || userId2 == uid'],
            ['ProGuard/code shrinking disabled', 'MEDIUM', 'Fix ML Kit compatibility, re-enable minification'],
            ['Catch-all rule exposes new collections', 'MEDIUM', 'Explicitly define rules for each new collection'],
            ['No rate limiting on callable functions', 'MEDIUM', 'Add App Check enforcement + per-user rate limits'],
        ]
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 10. BUILD & DEPLOYMENT
    # ══════════════════════════════════════════════════
    doc.add_heading('10. Build & Deployment', level=1)

    doc.add_heading('10.1 Android Build Configuration', level=2)
    doc.add_paragraph('File: android/app/build.gradle.kts')
    add_styled_table(doc,
        ['Setting', 'Value'],
        [
            ['compileSdk', 'flutter.compileSdkVersion (dynamic)'],
            ['minSdk', 'flutter.minSdkVersion (dynamic)'],
            ['targetSdk', 'flutter.targetSdkVersion (dynamic)'],
            ['NDK Version', '28.2.13676358'],
            ['Java/Kotlin', 'Java 17, Kotlin JVM target 17'],
            ['Core Library Desugaring', 'Enabled (Java 17 on older devices)'],
            ['Package ID', 'com.greengochat.greengochatapp'],
            ['Signing', 'key.properties (keystore, passwords, alias)'],
            ['Code Shrinking', 'DISABLED (ML Kit compatibility)'],
            ['ProGuard', 'NOT CONFIGURED'],
            ['Google Maps', 'API key from local.properties'],
        ]
    )

    doc.add_heading('10.2 iOS Build Configuration', level=2)
    add_styled_table(doc,
        ['Setting', 'Value'],
        [
            ['Minimum iOS', '15.5'],
            ['Framework Linking', 'Static frameworks with modular headers'],
            ['Bitcode', 'Disabled (deprecated Xcode 14+)'],
            ['C++ Standard', 'C++17 (gRPC/abseil compatibility)'],
            ['Script Sandboxing', 'Disabled (Xcode 15+ workaround)'],
            ['Background Modes', 'Fetch + Remote Notifications'],
            ['Permissions', 'Location, Camera, Microphone, Photo Library'],
            ['Google Maps API Key', 'Via Secrets.xcconfig'],
        ]
    )

    doc.add_heading('10.3 Firebase Deployment', level=2)
    add_code_block(doc, '''# Deploy commands:
firebase deploy                                    # Deploy everything
firebase deploy --only functions                   # All functions
firebase deploy --only functions:functionName      # Single function
firebase deploy --only firestore:rules             # Security rules
firebase deploy --only firestore:indexes           # Composite indexes
firebase deploy --only hosting:app                 # Flutter web
firebase deploy --only hosting:admin               # Admin panel

# CRITICAL: NEVER use --force for functions deployment
# It will DELETE all production functions not in local source

# Build commands:
flutter build apk --debug                          # Debug APK
flutter build apk --release                        # Release APK
flutter build appbundle --release                   # App Bundle (Play Store)
flutter build ios --release                        # iOS (requires Mac)''')

    doc.add_heading('10.4 Emulator Setup', level=2)
    add_styled_table(doc,
        ['Service', 'Host', 'Port'],
        [
            ['Auth', '0.0.0.0', '9099'],
            ['Functions', '0.0.0.0', '5001'],
            ['Firestore', '0.0.0.0', '8080'],
            ['Storage', '0.0.0.0', '9199'],
            ['Pub/Sub', '0.0.0.0', '8085'],
            ['Emulator UI', '0.0.0.0', '4000'],
        ]
    )
    doc.add_paragraph(
        'Start emulators: firebase emulators:start. '
        'App auto-detects emulator mode in kDebugMode and connects to 10.0.2.2 (Android) '
        'or 127.0.0.1 (iOS/Web). Auto-seeds 1000+ test profiles when < 1000 exist.'
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 11. LOCALIZATION
    # ══════════════════════════════════════════════════
    doc.add_heading('11. Localization Engineering', level=1)
    add_styled_table(doc,
        ['Setting', 'Value'],
        [
            ['ARB files location', 'lib/l10n/ (app_en.arb, app_de.arb, etc.)'],
            ['Generated output', 'lib/generated/ (NOT flutter_gen/gen_l10n/)'],
            ['Import path', 'package:greengo_chat/generated/app_localizations.dart'],
            ['Template', 'app_en.arb'],
            ['Languages', 'EN, DE, ES, FR, IT, PT, PT-BR (7 total)'],
            ['Config file', 'l10n.yaml at project root'],
        ]
    )
    doc.add_heading('Adding Translations', level=2)
    add_code_block(doc, '''// 1. Add key to app_en.arb:
"achievementRewardLabel": "{amount} {type} Reward",
"@achievementRewardLabel": {
  "placeholders": {
    "amount": { "type": "int" },
    "type": { "type": "String" }
  }
}

// 2. Add to all 6 other ARB files (de, es, fr, it, pt, pt_BR)
// 3. Run: flutter gen-l10n
// 4. Use: AppLocalizations.of(context)?.achievementRewardLabel(50, 'XP') ?? 'Fallback' ''')

    # ══════════════════════════════════════════════════
    # 12. APP INITIALIZATION
    # ══════════════════════════════════════════════════
    doc.add_heading('12. App Initialization Flow', level=1)
    doc.add_paragraph('File: lib/main.dart')
    numbered(doc, 'Lock orientation to portrait (portraitUp, portraitDown)')
    numbered(doc, 'Initialize Firebase with DefaultFirebaseOptions.currentPlatform')
    numbered(doc, 'Enable Firestore offline persistence with unlimited cache size')
    numbered(doc, 'If debug mode + emulator flag: connect to local emulators, disable Crashlytics/Performance, auto-seed test data')
    numbered(doc, 'Initialize Firebase App Check (Play Integrity on Android, App Attest on iOS)')
    numbered(doc, 'Configure Remote Config (10s fetch debug / 1hr production, default values for features)')
    numbered(doc, 'Initialize services: CacheService, FeatureFlagsService, AccessControlService, VersionCheckService, PushNotificationService, SoundService')
    numbered(doc, 'Register GetIt dependencies (injection_container.dart)')
    numbered(doc, 'Build MaterialApp with BlocProvider tree, localization delegates, named routes')
    numbered(doc, 'Initial route: LoginScreen -> MainNavigationScreen (post-auth)')

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 13. API REFERENCE
    # ══════════════════════════════════════════════════
    doc.add_heading('13. API Reference (Cloud Functions)', level=1)
    doc.add_paragraph(
        'All callable functions are invoked via Firebase SDK (CloudFunctions.instance.httpsCallable). '
        'Response format: { success: boolean, data?: T, message?: string }'
    )

    doc.add_heading('13.1 Gamification API', level=2)
    add_styled_table(doc,
        ['Function', 'Input', 'Output', 'Trigger Type'],
        [
            ['grantXP', '{action: string, metadata?: object}', '{success, xpAwarded, newLevel, totalXP}', 'onCall'],
            ['trackAchievementProgress', '{achievementId, progress}', '{success, isUnlocked}', 'onCall'],
            ['unlockAchievementReward', '{achievementId}', '{success, rewards}', 'onCall'],
            ['claimLevelRewards', '{level}', '{success, rewards}', 'onCall'],
            ['trackChallengeProgress', '{challengeId, progress}', '{success, isCompleted}', 'onCall'],
            ['claimChallengeReward', '{challengeId}', '{success, rewards}', 'onCall'],
            ['refreshMyStats', '(none)', '{success, stats: UserStats}', 'onCall'],
            ['onMessageCreatedVocabulary', '(auto)', 'void', 'onDocumentCreated'],
            ['computeDailyUserStats', '(auto)', 'void', 'onSchedule 3AM UTC'],
        ]
    )

    doc.add_heading('13.2 Coins API', level=2)
    add_styled_table(doc,
        ['Function', 'Input', 'Output', 'Notes'],
        [
            ['verifyGooglePlayCoinPurchase', '{purchaseToken, productId}', '{success, verified, coins}', 'Graceful if API unavailable'],
            ['verifyAppStoreCoinPurchase', '{receiptData, productId}', '{success, verified, coins}', 'Tries prod then sandbox'],
            ['claimReward', '{rewardId}', '{success, coinsAwarded}', 'Achievement/challenge rewards'],
        ]
    )

    doc.add_heading('13.3 Admin API', level=2)
    add_styled_table(doc,
        ['Function', 'Required Permission', 'Input', 'Notes'],
        [
            ['getUserActivityMetrics', 'viewDashboard', '(none)', 'DAU, WAU, MAU counts'],
            ['searchUsers', 'viewUserProfiles', '{query, filters, page}', 'Paginated, max 100/page'],
            ['suspendUserAccount', 'suspendUsers', '{userId, reason, duration}', 'Audit logged'],
            ['banUserAccount', 'banUsers', '{userId, reason}', 'Permanent, audit logged'],
            ['executeMassAction', 'executeMassActions', '{userIds[], action}', 'Bulk, max 1000 users'],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════
    # 14. DIAGRAMS
    # ══════════════════════════════════════════════════
    doc.add_heading('14. Diagrams Reference', level=1)
    doc.add_paragraph('Open with https://app.diagrams.net/ or any draw.io compatible editor.')
    add_styled_table(doc,
        ['File', 'Description'],
        [
            ['docs/diagrams/system_architecture.drawio', '4-layer architecture: Client, Firebase, Functions, External'],
            ['docs/diagrams/data_model.drawio', 'Firestore ERD with key collections and relationships'],
            ['docs/diagrams/discovery_flow.drawio', 'Discovery pipeline: parallel queries, scoring, matching'],
            ['docs/diagrams/gamification_flow.drawio', 'XP actions, level thresholds, vocabulary, celebrations'],
            ['docs/diagrams/message_lifecycle.drawio', 'Message: send -> triggers -> push/translate/vocab -> read'],
        ]
    )

    doc.add_paragraph()
    doc.add_heading('Document Information', level=2)
    add_styled_table(doc,
        ['Item', 'Value'],
        [
            ['Version', '2.0 (Engineering Edition)'],
            ['Date', 'March 10, 2026'],
            ['App Version', '1.1.0+26'],
            ['Firebase Project', 'greengo-chat'],
            ['Package', 'com.greengochat.greengochatapp'],
            ['Flutter SDK', '>=3.0.0 <4.0.0'],
            ['Node.js Runtime', '22'],
            ['Classification', 'Confidential - Engineering Team Only'],
        ]
    )

    return doc


if __name__ == '__main__':
    doc = create_document()
    output_path = os.path.join(os.path.dirname(__file__), 'GreenGo_Technical_Documentation.docx')
    doc.save(output_path)
    print(f'Documentation generated: {output_path}')
    # Also report size
    size = os.path.getsize(output_path)
    print(f'File size: {size / 1024:.1f} KB')
